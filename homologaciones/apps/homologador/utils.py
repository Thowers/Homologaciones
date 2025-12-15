import PyPDF2
import io
from django.core.files.uploadedfile import InMemoryUploadedFile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import json

def extraer_texto_de_archivo(archivo: InMemoryUploadedFile) -> str:
    """Extrae texto de un archivo PDF subido en Django."""
    
    # Es crucial que el archivo se lea solo una vez y en modo binario
    try:
        archivo_bytes = archivo.read()
    except Exception as e:
        return f"Error al leer archivo binario: {e}"
    
    
    if archivo.content_type == 'application/pdf':
        try:
            lector = PyPDF2.PdfReader(io.BytesIO(archivo_bytes))
            texto_completo = ""
            
            for pagina in lector.pages:
                # 🚨 CORRECCIÓN CLAVE: 🚨
                # 1. Extraer el texto de la página.
                # 2. Concatenarlo a texto_completo, añadiendo un salto de línea (espacio) para separar páginas.
                texto_completo += pagina.extract_text() + "\n"
                
            # Opcional: limpiar la cadena final de espacios extra
            return texto_completo.strip()
            
        except Exception as e:
            # Capturar errores específicos de PyPDF2 (ej. PDF encriptado)
            return f"Error al procesar PDF (PyPDF2): {e}"
    
    else:
        return "Tipo de archivo no soportado (solo PDF por ahora)."
    
def generar_docx_homologacion(historico_obj):
    """
    Genera un documento DOCX con el resultado de la homologación.
    Toma un objeto HistoricoHomologacion.
    """
    document = Document()

    # --- Estilos de fuente UMB (Simulación) ---
    style_heading = document.styles['Heading 1']
    style_heading.font.name = 'Arial'
    style_heading.font.size = Pt(16)
    style_heading.font.color.rgb = RGBColor(0x00, 0x33, 0x66) # Azul UMB

    style_normal = document.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)

    # --- Título ---
    document.add_heading('Informe de Homologación de Asignaturas', 0)
    
    # --- Metadatos ---
    document.add_paragraph(f'Carrera de Destino: {historico_obj.carrera_destino}')
    document.add_paragraph(f'Estudiante (ID): {historico_obj.documento_identidad or "N/A"}')
    document.add_paragraph(f'Fecha de Proceso: {historico_obj.fecha_procesamiento.strftime("%d/%m/%Y %H:%M:%S")}')
    document.add_paragraph().add_run().add_break()
    
    # --- Contenido de la Homologación ---
    
    resultados = historico_obj.resultado_parsed
    if not resultados:
        document.add_paragraph('No se encontraron resultados válidos para esta homologación.')
        
        # Guardar el documento en memoria
        f = io.BytesIO()
        document.save(f)
        f.seek(0)
        return f

    # Dividir resultados
    homologadas = [res for res in resultados if res.get('estado') == 'HOMOLOGADA']
    no_aplica = [res for res in resultados if res.get('estado') == 'NO APLICA']
    
    # ----------------------------------------------------
    # TABLA DE HOMOLOGADAS
    # ----------------------------------------------------
    document.add_heading('Asignaturas Homologadas', 1)
    
    if homologadas:
        table = document.add_table(rows=1, cols=4, style='Light Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Materia Destino'
        hdr_cells[1].text = 'Materia Origen'
        hdr_cells[2].text = 'Créditos Otorgados'
        hdr_cells[3].text = 'Razón'
        
        for item in homologadas:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{item['materia_destino']} ({item['codigo_destino']})"
            row_cells[1].text = item['materia_origen_homologada']
            row_cells[2].text = str(item['creditos_otorgados'])
            row_cells[3].text = item['razon_homologacion']
    else:
        document.add_paragraph('Ninguna asignatura pudo ser homologada.')
        
    # ----------------------------------------------------
    # SECCIÓN NO APLICA
    # ----------------------------------------------------
    document.add_paragraph().add_run().add_break()
    document.add_heading('Asignaturas No Homologadas (No Aplica)', 1)

    if no_aplica:
        for item in no_aplica:
            p = document.add_paragraph()
            p.add_run(f"• {item['materia_destino']} ({item['codigo_destino']}): ").bold = True
            p.add_run(f"{item['razon_homologacion']}")
    else:
        document.add_paragraph('Todas las asignaturas pudieron ser homologadas o la lista de destino era vacía.')


    # Guardar el documento en memoria (no en el disco)
    f = io.BytesIO()
    document.save(f)
    f.seek(0)
    return f